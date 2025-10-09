# app.py
import streamlit as st
from docx import Document
import io
import os
import re
import mammoth
import requests
from datetime import datetime
try:
    # Python 3.9+ zoneinfo
    from zoneinfo import ZoneInfo
    KOLKATA = ZoneInfo("Asia/Kolkata")
except Exception:
    # fallback
    import pytz
    KOLKATA = pytz.timezone("Asia/Kolkata")

st.set_page_config(page_title="ModFar COA Generator", layout="wide")
st.title("📄 ModFar COA Generator")

# --- Defaults: local paths (these match the files you uploaded) ---
DEFAULT_TEMPLATES = {
    "MOD": "PH LIPL MOD COA.docx",
    "FAR": "PH LIPL FAR COA.docx"
}

# --- Regex for placeholders like {{KEY}} ---
PLACEHOLDER_RE = re.compile(r"\{\{\s*([A-Za-z0-9_\-/]+)\s*\}\}")

# --- Utility to normalize common broken placeholders e.g. '((BATCH_1}}' -> '{{BATCH_1}}' ---
def normalize_broken_placeholders_in_doc(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            if "((" in run.text or "))" in run.text:
                run.text = run.text.replace("((", "{{").replace("))", "}}")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "((" in run.text or "))" in run.text:
                            run.text = run.text.replace("((", "{{").replace("))", "}}")
    try:
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for run in para.runs:
                    if "((" in run.text or "))" in run.text:
                        run.text = run.text.replace("((", "{{").replace("))", "}}")
            footer = section.footer
            for para in footer.paragraphs:
                for run in para.runs:
                    if "((" in run.text or "))" in run.text:
                        run.text = run.text.replace("((", "{{").replace("))", "}}")
    except Exception:
        pass

# --- Replace placeholders in a paragraph while preserving style of first overlapping run ---
def replace_placeholders_in_paragraph(paragraph, replacements):
    runs = paragraph.runs
    if not runs:
        return

    # Build full text and offsets
    full_text = ""
    offsets = []  # (run, start, end)
    for run in runs:
        start = len(full_text)
        full_text += run.text
        end = len(full_text)
        offsets.append((run, start, end))

    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    for match in matches:
        key = match.group(1)
        if key not in replacements:
            continue
        replacement_text = str(replacements[key])
        p_start, p_end = match.start(), match.end()

        overlapping = [(r, s, e) for (r, s, e) in offsets if not (e <= p_start or s >= p_end)]
        if not overlapping:
            continue

        first_run, first_s, first_e = overlapping[0]
        last_run, last_s, last_e = overlapping[-1]

        # prefix & suffix around placeholder inside first/last runs
        prefix_len = max(0, p_start - first_s)
        prefix = first_run.text[:prefix_len]
        suffix_start_in_last = p_end - last_s
        suffix = last_run.text[suffix_start_in_last:]

        # clear overlapping runs
        for r, _, _ in overlapping:
            r.text = ""

        new_text = prefix + replacement_text + suffix
        first_run.text = new_text

        # copy style from the first overlapping run (safe copy)
        try:
            font = first_run.font
            # assign only if attributes exist (some may be None)
            # copying same run's font to itself is harmless; copying from a style_run would be similar
            # (we keep this for compatibility if you prefer to choose a particular run)
            # note: if you wanted, pick 'style_run' differently
            if font.name:
                first_run.font.name = font.name
            if font.size:
                first_run.font.size = font.size
            first_run.font.bold = font.bold
            first_run.font.italic = font.italic
            first_run.font.underline = font.underline
            if font.color and getattr(font.color, "rgb", None) is not None:
                first_run.font.color.rgb = font.color.rgb
        except Exception:
            pass

def advanced_replace_text_preserving_style(doc, replacements):
    normalize_broken_placeholders_in_doc(doc)

    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, replacements)

    # headers/footers
    try:
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                replace_placeholders_in_paragraph(para, replacements)
            footer = section.footer
            for para in footer.paragraphs:
                replace_placeholders_in_paragraph(para, replacements)
    except Exception:
        pass

# --- Convert to HTML for preview using mammoth ---
def docx_to_html_bytes(docx_bytes):
    result = mammoth.convert_to_html(io.BytesIO(docx_bytes))
    return result.value

# --- Helper: download raw file from GitHub raw URL (returns bytes or None) ---
def download_from_github_raw(raw_url):
    try:
        r = requests.get(raw_url, timeout=15)
        r.raise_for_status()
        return r.content
    except Exception as e:
        st.warning(f"Failed to download from GitHub URL: {e}")
        return None

# --- Sidebar: template source options ---
st.sidebar.header("Template Source / Options")
use_default_paths = st.sidebar.checkbox("Use server default templates (if available)", value=True)
st.sidebar.markdown("You can also upload templates or provide GitHub raw URLs.")

uploaded_mod = None
uploaded_far = None
mod_github_url = st.sidebar.text_input("MOD template raw GitHub URL (optional)", value="")
far_github_url = st.sidebar.text_input("FAR template raw GitHub URL (optional)", value="")

file_source = st.sidebar.selectbox("If uploading, choose which template to upload", ["None", "Upload MOD", "Upload FAR", "Upload Both"])
if file_source in ("Upload MOD", "Upload Both"):
    uploaded_mod = st.sidebar.file_uploader("Upload MOD .docx", type=["docx"], key="upload_mod")
if file_source in ("Upload FAR", "Upload Both"):
    uploaded_far = st.sidebar.file_uploader("Upload FAR .docx", type=["docx"], key="upload_far")

# Which COA to generate
coa_type = st.selectbox("Choose COA type", ["MOD", "FAR"])

# Determine template bytes for selected COA
def get_template_bytes(coa_type):
    # priority: uploaded in sidebar -> GitHub raw URL -> default local path
    if coa_type == "MOD" and uploaded_mod is not None:
        return uploaded_mod.read()
    if coa_type == "FAR" and uploaded_far is not None:
        return uploaded_far.read()

    if coa_type == "MOD" and mod_github_url.strip():
        b = download_from_github_raw(mod_github_url.strip())
        if b:
            return b
    if coa_type == "FAR" and far_github_url.strip():
        b = download_from_github_raw(far_github_url.strip())
        if b:
            return b

    # fallback: default file path
    path = DEFAULT_TEMPLATES.get(coa_type)
    if path and os.path.exists(path):
        with open(path, "rb") as f:
            return f.read()
    return None

# --- Default date (DDMMYYYY) in Asia/Kolkata timezone, editable by user ---
now_kolkata = datetime.now(KOLKATA)
default_ddmmyyyy = now_kolkata.strftime("%d%m%Y")

st.markdown("### Enter values (four batches)")
with st.form("coa_form"):
    date_field = st.text_input("Date (DDMMYYYY)", value=default_ddmmyyyy)

    # Common batch labels
    st.write("#### Batch labels")
    batch_1 = st.text_input("BATCH_1", value="")
    batch_2 = st.text_input("BATCH_2", value="")
    batch_3 = st.text_input("BATCH_3", value="")
    batch_4 = st.text_input("BATCH_4", value="")

    # Fields differ by COA type
    if coa_type == "MOD":
        st.write("#### MOD fields (Moisture, 30min, 60min, pH)")
        # Moisture
        m1 = st.text_input("M1 (Batch1 Moisture)", value="")
        m2 = st.text_input("M2 (Batch2 Moisture)", value="")
        m3 = st.text_input("M3 (Batch3 Moisture)", value="")
        m4 = st.text_input("M4 (Batch4 Moisture)", value="")

        # Viscosities: 30min -> B1V1..B4V1 ; 60min -> B1V2..B4V2
        b1v1 = st.text_input("B1V1 (Batch1 - 30min viscosity)", value="")
        b2v1 = st.text_input("B2V1 (Batch2 - 30min viscosity)", value="")
        b3v1 = st.text_input("B3V1 (Batch3 - 30min viscosity)", value="")
        b4v1 = st.text_input("B4V1 (Batch4 - 30min viscosity)", value="")

        b1v2 = st.text_input("B1V2 (Batch1 - 60min viscosity)", value="")
        b2v2 = st.text_input("B2V2 (Batch2 - 60min viscosity)", value="")
        b3v2 = st.text_input("B3V2 (Batch3 - 60min viscosity)", value="")
        b4v2 = st.text_input("B4V2 (Batch4 - 60min viscosity)", value="")

        # pH
        ph1 = st.text_input("PH1 (Batch1 pH)", value="")
        ph2 = st.text_input("PH2 (Batch2 pH)", value="")
        ph3 = st.text_input("PH3 (Batch3 pH)", value="")
        ph4 = st.text_input("PH4 (Batch4 pH)", value="")

    else:  # FAR
        st.write("#### FAR fields (Moisture, 2h, 24h, pH, mesh, BD, Fann, Fann30)")
        # Moisture
        m1 = st.text_input("M1 (Batch1 Moisture)", value="")
        m2 = st.text_input("M2 (Batch2 Moisture)", value="")
        m3 = st.text_input("M3 (Batch3 Moisture)", value="")
        m4 = st.text_input("M4 (Batch4 Moisture)", value="")

        # Viscosities: 2h -> B1V1.. ; 24h -> B1V2..
        b1v1 = st.text_input("B1V1 (Batch1 - 2h viscosity)", value="")
        b2v1 = st.text_input("B2V1 (Batch2 - 2h viscosity)", value="")
        b3v1 = st.text_input("B3V1 (Batch3 - 2h viscosity)", value="")
        b4v1 = st.text_input("B4V1 (Batch4 - 2h viscosity)", value="")

        b1v2 = st.text_input("B1V2 (Batch1 - 24h viscosity)", value="")
        b2v2 = st.text_input("B2V2 (Batch2 - 24h viscosity)", value="")
        b3v2 = st.text_input("B3V2 (Batch3 - 24h viscosity)", value="")
        b4v2 = st.text_input("B4V2 (Batch4 - 24h viscosity)", value="")

        # pH
        ph1 = st.text_input("PH1 (Batch1 pH)", value="")
        ph2 = st.text_input("PH2 (Batch2 pH)", value="")
        ph3 = st.text_input("PH3 (Batch3 pH)", value="")
        ph4 = st.text_input("PH4 (Batch4 pH)", value="")

        # Mesh
        mesh1 = st.text_input("MESH1 (Batch1 through 200 mesh %)", value="")
        mesh2 = st.text_input("MESH2 (Batch2 through 200 mesh %)", value="")
        mesh3 = st.text_input("MESH3 (Batch3 through 200 mesh %)", value="")
        mesh4 = st.text_input("MESH4 (Batch4 through 200 mesh %)", value="")

        # Bulk Density
        bd1 = st.text_input("BD1 (Batch1 Bulk Density)", value="")
        bd2 = st.text_input("BD2 (Batch2 Bulk Density)", value="")
        bd3 = st.text_input("BD3 (Batch3 Bulk Density)", value="")
        bd4 = st.text_input("BD4 (Batch4 Bulk Density)", value="")

        # Fann 3'
        f1 = st.text_input("F1 (Batch1 Fann @ 3')", value="")
        f2 = st.text_input("F2 (Batch2 Fann @ 3')", value="")
        f3 = st.text_input("F3 (Batch3 Fann @ 3')", value="")
        f4 = st.text_input("F4 (Batch4 Fann @ 3')", value="")

        # Fann 30'
        fv1 = st.text_input("FV1 (Batch1 Fann @ 30')", value="")
        fv2 = st.text_input("FV2 (Batch2 Fann @ 30')", value="")
        fv3 = st.text_input("FV3 (Batch3 Fann @ 30')", value="")
        fv4 = st.text_input("FV4 (Batch4 Fann @ 30')", value="")

    submitted = st.form_submit_button("Generate COA")

# --- On submit: build replacements and process template ---
if submitted:
    template_bytes = get_template_bytes(coa_type)
    if template_bytes is None:
        st.error("Template not available. Provide via default path, upload, or GitHub raw URL.")
    else:
        # load docx from bytes into python-docx Document
        try:
            doc = Document(io.BytesIO(template_bytes))
        except Exception as e:
            st.error(f"Failed to open template as docx: {e}")
            doc = None

        if doc:
            # Build replacements based on COA type
            replacements = {}
            # Date mapping: two templates use different date placeholder styles in your files:
            # - MOD file uses {{DD/MM/YYYY}} (per inspection)
            # - FAR file uses {{DD-MM-YYYY}}
            # We'll populate both keys so whichever exists in template will be replaced.
            replacements["DD/MM/YYYY"] = date_field
            replacements["DD-MM-YYYY"] = date_field

            # Batch labels
            replacements.update({
                "BATCH_1": batch_1,
                "BATCH_2": batch_2,
                "BATCH_3": batch_3,
                "BATCH_4": batch_4
            })

            # Common moisture & viscosities & pH
            replacements.update({
                "M1": m1, "M2": m2, "M3": m3, "M4": m4,
                "B1V1": b1v1, "B2V1": b2v1, "B3V1": b3v1, "B4V1": b4v1,
                "B1V2": b1v2, "B2V2": b2v2, "B3V2": b3v2, "B4V2": b4v2,
                "PH1": ph1, "PH2": ph2, "PH3": ph3, "PH4": ph4
            })

            # Additional FAR-only fields
            if coa_type == "FAR":
                replacements.update({
                    "MESH1": mesh1, "MESH2": mesh2, "MESH3": mesh3, "MESH4": mesh4,
                    "BD1": bd1, "BD2": bd2, "BD3": bd3, "BD4": bd4,
                    "F1": f1, "F2": f2, "F3": f3, "F4": f4,
                    "FV1": fv1, "FV2": fv2, "FV3": fv3, "FV4": fv4
                })

            # Perform replacement
            advanced_replace_text_preserving_style(doc, replacements)

            # Save to BytesIO
            out_buffer = io.BytesIO()
            doc.save(out_buffer)
            out_bytes = out_buffer.getvalue()

            # Preview using mammoth (HTML)
            try:
                html = docx_to_html_bytes(out_bytes)
                st.subheader("📄 Preview (HTML)")
                st.components.v1.html(f"<div style='padding:12px'>{html}</div>", height=700, scrolling=True)
            except Exception as e:
                st.warning(f"Preview (mammoth) failed: {e}. You can still download the DOCX.")

            # Download button
            filename = f"COA_{coa_type}_{batch_1 or 'batch'}.docx"
            st.download_button(
                label="📥 Download generated DOCX",
                data=out_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            st.success("Generated. Open the downloaded DOCX in MS Word to confirm visual formatting.")
